from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


PROJECT_ROOT = Path(__file__).resolve().parents[1]
REPORT_PATH = PROJECT_ROOT / "report.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def add_link(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(color)
    r_pr.append(underline)
    run.append(r_pr)
    t = OxmlElement("w:t")
    t.text = text
    run.append(t)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def configure_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def add_title(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run("Persistent Identity Memory for Long-Horizon ConsisID Generation")
    run.font.name = "Calibri"
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor.from_string("0B2545")
    run.bold = True
    sub = doc.add_paragraph()
    sub.add_run(f"Living literature and implementation report - updated {date.today().isoformat()}").italic = True


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_key_value_table(doc, rows):
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    widths = [2200, 7160]
    for index, heading in enumerate(["Area", "Current map"]):
        cell = table.rows[0].cells[index]
        cell.text = heading
        set_cell_shading(cell, "F2F4F7")
        set_cell_width(cell, widths[index])
    for key, value in rows:
        row = table.add_row()
        row.cells[0].text = key
        row.cells[1].text = value
        for idx, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            set_cell_width(cell, widths[idx])
    return table


def add_literature_table(doc, papers):
    table = doc.add_table(rows=1, cols=7)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    headers = ["Paper", "Problem", "Method", "Strengths", "Limits", "Project relevance", "Transferable ideas"]
    widths = [1500, 1450, 1650, 1350, 1350, 1450, 1600]
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        set_cell_shading(cell, "F2F4F7")
        set_cell_width(cell, widths[i])
    for paper in papers:
        row = table.add_row()
        values = [
            paper["paper"],
            paper["problem"],
            paper["method"],
            paper["strengths"],
            paper["limits"],
            paper["relevance"],
            paper["transfer"],
        ]
        for i, value in enumerate(values):
            row.cells[i].text = value
            row.cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            set_cell_width(row.cells[i], widths[i])
    return table


def add_paper_notes(doc, papers):
    for paper in papers:
        doc.add_heading(paper["paper"], level=2)
        for label, key in [
            ("Problem", "problem"),
            ("Method", "method"),
            ("Strengths", "strengths"),
            ("Limitations", "limits"),
            ("Relevance to this project", "relevance"),
            ("Transferable ideas", "transfer"),
        ]:
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(f"{label}: ")
            run.bold = True
            paragraph.add_run(paper[key])


def build():
    doc = Document()
    configure_document(doc)
    add_title(doc)

    doc.add_heading("Executive Position", level=1)
    doc.add_paragraph(
        "The current main architecture remains the right near-term direction: ConsisID-style frequency-aware identity "
        "conditioning plus an identity-specialized memory layer. The codebase already exposes identity tensors before "
        "generation, so a retrieval prototype can be added without rewriting the DiT or merging large external repos."
    )
    add_bullets(
        doc,
        [
            "Main drift hypothesis still looks strong: most 2025-2026 long-video/world-model papers frame failure as finite-context forgetting, error accumulation, or missing memory retrieval.",
            "The first implementation should stay outside the transformer: retrieve high-confidence identity episodes, blend identity conditioning tensors, then measure ArcFace decay by chunk.",
            "The next deeper insertion point is transformer-local: recurrent identity state or time-aware identity injection around the local facial extractor and identity cross-attention.",
            "Do not replace ConsisID yet. Newer methods suggest useful modules, but none obviously removes the need for persistent identity memory.",
        ],
    )

    doc.add_heading("Codebase Map", level=1)
    add_key_value_table(
        doc,
        [
            (
                "Inference pipeline",
                "infer.py validates a PNG identity image or loads a saved identity memory, prepares face models, loads ConsisIDTransformer3DModel and ConsisIDPipeline, extracts or reloads identity tensors, calls pipe(...), then exports frames to MP4.",
            ),
            (
                "Identity extraction",
                "models/consisid_utils.py uses InsightFace AntelopeV2 for a 512-d ArcFace-like embedding, facexlib alignment/parsing for face crop isolation, and EVA02-CLIP-L visual features. These form id_cond = [id_ante_embedding, id_cond_vit] plus id_vit_hidden.",
            ),
            (
                "Identity injection",
                "models/pipeline_consisid.py forwards id_cond and id_vit_hidden to the transformer at every denoising step. kps_cond is optionally drawn and VAE-encoded as extra image latents.",
            ),
            (
                "Transformer path",
                "models/transformer_consisid.py turns id_cond and id_vit_hidden into valid_face_emb through local_facial_extractor, then injects that face embedding through PerceiverCrossAttention every cross_attn_interval transformer blocks.",
            ),
            (
                "Best memory insertion points",
                "1. Pre-pipeline identity tensor retrieval/blending in infer.py. 2. Pipeline-level scheduling of identity strength by denoising step. 3. Transformer-level recurrent identity state around valid_face_emb. 4. Future training-time memory losses.",
            ),
            (
                "Evaluation",
                "eval/arcface_identity_stability.py computes per-frame and per-chunk InsightFace similarity, detection rate, min/mean/max similarity, lowest chunk, and now chunk-wise decay/slope.",
            ),
        ],
    )

    doc.add_heading("Implemented First Prototype", level=1)
    add_bullets(
        doc,
        [
            "Added util/identity_memory.py for episodic identity bank loading, retrieval, and blending.",
            "Added infer.py flags: --episodic_identity_memory_path, --episodic_top_k, --episodic_min_similarity, --episodic_base_weight, --episodic_update_memory, and --episodic_memory_max_episodes.",
            "Added tools/identity_memory_bank.py to convert an existing single identity_memory.pt into a reusable episodic bank.",
            "Added eval/identity_prompt_suite.json and eval/run_identity_persistence_experiment.py for baseline vs episodic experiments over realistic and stylized prompt categories.",
        ],
    )

    doc.add_heading("Evaluation Protocol", level=1)
    add_bullets(
        doc,
        [
            "Primary metric: ArcFace/InsightFace cosine similarity to the reference identity embedding.",
            "Chunk metric: mean/min/max similarity per 49-frame chunk, with face detection rate tracked separately so occlusion is not mistaken for identity drift.",
            "Decay metric: last chunk minus first chunk and linear slope of chunk mean similarity.",
            "Prompt split: realistic prompts for natural face preservation and stylized prompts for anime, manga, game-like, and painterly identity transfer.",
            "Comparisons: baseline, episodic memory, recurrent state, and episodic plus recurrent once the recurrent module exists.",
        ],
    )

    doc.add_heading("Literature Review", level=1)
    doc.add_paragraph(
        "Chrome/Google Scholar automation was attempted, but the local Chrome extension sandbox failed to initialize. "
        "This first sweep therefore uses primary arXiv, CVF, and publisher/research-lab pages. Re-run Google Scholar "
        "later for citation counts and related-paper graph expansion."
    )

    papers = [
        {
            "paper": "ConsisID, CVPR 2025",
            "problem": "Tuning-free identity-preserving text-to-video from one image.",
            "method": "Frequency-aware DiT control: low-frequency global face features plus high-frequency intrinsic identity features.",
            "strengths": "Directly matches this codebase; strong local identity preservation.",
            "limits": "No explicit long-term memory; mostly short-clip conditioning.",
            "relevance": "Baseline and host architecture.",
            "transfer": "Keep face-frequency decomposition; add memory outside first.",
        },
        {
            "paper": "HunyuanCustom, 2025",
            "problem": "Customized video generation from multimodal conditions.",
            "method": "LLaVA-based image-text fusion and image-ID enhancement through temporal concatenation.",
            "strengths": "Shows identity reinforcement can be framed as repeated temporal conditioning.",
            "limits": "Not specialized for identity drift over long autoregressive rollouts.",
            "relevance": "Supports identity propagation direction.",
            "transfer": "Use repeated retrieved identity features as temporal anchors.",
        },
        {
            "paper": "Proteus-ID, 2025",
            "problem": "Identity consistency and natural motion in video customization.",
            "method": "Multimodal identity fusion, time-aware identity injection, adaptive motion learning.",
            "strengths": "Explicitly balances identity and motion.",
            "limits": "Requires trained modules and dataset; not a minimal drop-in.",
            "relevance": "Strong candidate for second-stage injection schedule.",
            "transfer": "Time-aware identity strength; motion-aware loss weighting.",
        },
        {
            "paper": "LaVieID, 2025",
            "problem": "Stochastic DiT generation loses local facial identity across time.",
            "method": "Local facial router plus temporal autoregressive latent refinement.",
            "strengths": "Direct evidence for local plus temporal identity modeling.",
            "limits": "Architectural change, likely not first prototype.",
            "relevance": "Closest to RAD-style identity state among ID papers.",
            "transfer": "Recurrent bias/refinement over face-related latent tokens.",
        },
        {
            "paper": "MagicID, ICCV 2025",
            "problem": "Identity preservation often reduces motion; longer videos degrade ID.",
            "method": "Hybrid preference optimization with identity and dynamics reward pairs.",
            "strengths": "Treats identity-motion trade-off as optimization target.",
            "limits": "Preference data and training required.",
            "relevance": "Useful for later training/evaluation.",
            "transfer": "Construct pairwise identity decay preferences from our evaluator.",
        },
        {
            "paper": "ContextAnyone, 2025",
            "problem": "Reference-to-video methods preserve face but lose hair, outfit, body shape.",
            "method": "Joint reference reconstruction and video generation, Emphasize-Attention, Gap-RoPE.",
            "strengths": "Broad character consistency, not just face.",
            "limits": "New DiT modifications; late-stage replacement candidate only.",
            "relevance": "Important if identity includes non-face cues.",
            "transfer": "Separate reference/memory token positions; reinforce reference-aware features.",
        },
        {
            "paper": "WorldMem, 2025/2026",
            "problem": "World simulators forget scenes beyond context windows.",
            "method": "Memory bank storing memory frames and states such as pose and timestamp; memory attention retrieves relevant units.",
            "strengths": "Clean memory-bank formulation for long consistency.",
            "limits": "Scene/spatial memory, not identity-specialized.",
            "relevance": "Primary inspiration for episodic retrieval.",
            "transfer": "Store identity episodes with state: prompt, style, time, confidence, viewpoint.",
        },
        {
            "paper": "RAD, 2025",
            "problem": "Autoregressive video diffusion lacks memory compression/retrieval beyond window size.",
            "method": "RNN/LSTM memory integrated with diffusion transformer; frame-wise autoregression for update and retrieval.",
            "strengths": "Matches recurrent identity-state hypothesis.",
            "limits": "Requires training/inference alignment; not a pure inference hack.",
            "relevance": "Second milestone after episodic memory.",
            "transfer": "Maintain recurrent identity state over chunks; compare against retrieval bank.",
        },
        {
            "paper": "Video World Models with Long-term Spatial Memory, 2025",
            "problem": "World models forget revisited environments.",
            "method": "Geometry-grounded long-term spatial memory with store/retrieve mechanisms.",
            "strengths": "Explicit long-term memory evaluation.",
            "limits": "Geometry-focused; identity needs face-aware memory state.",
            "relevance": "Supports specialized memory rather than longer context only.",
            "transfer": "Evaluate revisits/occlusion as memory tests, not only continuous similarity.",
        },
        {
            "paper": "WorldPack, 2025",
            "problem": "Long context is too expensive for spatially consistent world rollouts.",
            "method": "Trajectory packing plus memory retrieval.",
            "strengths": "Shows compression plus retrieval can beat brute context.",
            "limits": "World navigation domain.",
            "relevance": "Good argument for compact identity memory bank.",
            "transfer": "Compress identity episodes by representative embeddings rather than storing all frames.",
        },
        {
            "paper": "LIVE, 2026",
            "problem": "Autoregressive rollout errors accumulate beyond training horizon.",
            "method": "Cycle-consistency objective and progressive training to bound error.",
            "strengths": "Direct long-horizon error-control framing.",
            "limits": "Training objective, not immediate inference module.",
            "relevance": "Explains why local identity may decay despite strong first chunk.",
            "transfer": "Later train identity-cycle consistency: recover initial ID after long rollout.",
        },
        {
            "paper": "Identity-Motion Trade-offs, BMVC 2025",
            "problem": "Motion, structure, and identity entangle in T2V attention features.",
            "method": "Analysis of self-attention Q features; Q-injection for motion transfer and multi-shot consistency.",
            "strengths": "Mechanistic clue for why identity changes when motion/viewpoint changes.",
            "limits": "Training-free control, not identity memory by itself.",
            "relevance": "Useful for understanding drift causes.",
            "transfer": "Probe attention/query features during drift; avoid memory injection that suppresses motion.",
        },
        {
            "paper": "Identity-GRPO, 2025",
            "problem": "Multi-human interactions cause identity swaps/degradation.",
            "method": "Human-consistency reward model plus GRPO optimization.",
            "strengths": "Gives reward-learning route for identity consistency.",
            "limits": "RL training overhead; multi-subject focus.",
            "relevance": "Later alignment stage if inference memory helps but is not enough.",
            "transfer": "Use evaluator outputs to build automatic reward signals.",
        },
        {
            "paper": "DreamVideo-Omni, 2026",
            "problem": "Multi-subject identity and multi-granularity motion control degrade together.",
            "method": "Condition-aware 3D RoPE, hierarchical motion injection, role embeddings, latent identity reward learning.",
            "strengths": "Modern 2026 evidence for identity-aware rewards and role binding.",
            "limits": "Large training framework.",
            "relevance": "Useful if project expands to multiple characters.",
            "transfer": "Role embeddings for subject-specific identity memory slots.",
        },
        {
            "paper": "LumosX, 2026",
            "problem": "Face-attribute alignment fails in personalized multi-subject video.",
            "method": "Relational self/cross attention with face-attribute dependencies.",
            "strengths": "Makes identity broader than face embedding.",
            "limits": "Multi-subject and dataset-heavy.",
            "relevance": "Stylized identity needs hair, outfit, silhouette, and accessories.",
            "transfer": "Extend identity memory from face-only to face plus attribute slots.",
        },
    ]
    add_paper_notes(doc, papers)

    doc.add_heading("Source Links", level=1)
    sources = [
        ("ConsisID CVPR 2025 poster", "https://cvpr.thecvf.com/virtual/2025/poster/32871"),
        ("ConsisID paper PDF", "https://openaccess.thecvf.com/content/CVPR2025/papers/Yuan_Identity-Preserving_Text-to-Video_Generation_by_Frequency_Decomposition_CVPR_2025_paper.pdf"),
        ("HunyuanCustom", "https://arxiv.org/abs/2505.04512"),
        ("Proteus-ID", "https://arxiv.org/abs/2506.23729"),
        ("LaVieID", "https://arxiv.org/abs/2508.07603"),
        ("MagicID", "https://arxiv.org/abs/2503.12689"),
        ("ContextAnyone", "https://arxiv.org/abs/2512.07328"),
        ("WorldMem", "https://arxiv.org/abs/2504.12369"),
        ("RAD", "https://arxiv.org/abs/2511.12940"),
        ("Video World Models with Long-term Spatial Memory", "https://arxiv.org/abs/2506.05284"),
        ("WorldPack", "https://arxiv.org/abs/2512.02473"),
        ("LIVE", "https://arxiv.org/abs/2602.03747"),
        ("Identity-Motion Trade-offs", "https://research.nvidia.com/publication/2025-07_identity-motion-trade-offs-text-video-generation"),
        ("Identity-GRPO", "https://arxiv.org/abs/2510.14256"),
        ("DreamVideo-Omni", "https://arxiv.org/abs/2603.12257"),
        ("LumosX", "https://arxiv.org/abs/2603.20192"),
    ]
    for title, url in sources:
        p = doc.add_paragraph(style="List Bullet")
        add_link(p, title, url)

    doc.add_heading("Next Experimental Decision", level=1)
    doc.add_paragraph(
        "Run the new suite on the cluster with baseline and episodic memory. If episodic retrieval improves late-chunk "
        "similarity or reduces negative slope without flattening motion, keep the architecture and move to recurrent "
        "identity state. If it only improves first-frame similarity or over-constrains stylized prompts, shift toward "
        "time-aware injection and attention-level identity scheduling before adding recurrent training."
    )

    doc.save(REPORT_PATH)
    print(f"Wrote {REPORT_PATH}")


if __name__ == "__main__":
    build()
